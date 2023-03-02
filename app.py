import docx
import os
from docx.oxml import OxmlElement #Para caption (referencias de tabla)
from docx.oxml.ns import qn #Para caption (referencias de tabla)
from docx.enum.text import WD_ALIGN_PARAGRAPH

#Funciones para agregar caption
def Figure(paragraph):
    run = run = paragraph.add_run()
    r = run._r
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    r.append(fldChar)
    instrText = OxmlElement('w:instrText')
    instrText.text = ' SEQ Figure * ARABIC' #Aqui cambia
    r.append(instrText)
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    r.append(fldChar)
def Table(paragraph):
    run = run = paragraph.add_run()
    r = run._r
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    r.append(fldChar)
    instrText = OxmlElement('w:instrText')
    instrText.text = ' SEQ Table * ARABIC' #Aqui cambia
    r.append(instrText)
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    r.append(fldChar)
def Referencia(paragraph, grafico=False):
    #Función para crear una referencia a una tabla o un gráfico
    p = doc.add_paragraph()
    p.add_run("\n"+paragraph).italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if grafico == True:
        Figure(p)
    else:
        Table(p)

def Indice():
    # Agregar un índice al final del documento
    doc.add_page_break()
    indice = doc.add_paragraph()
    indice.add_run('Índice\n').bold = True
    for p in doc.paragraphs:
        if 'Heading' in p.style.name:
            texto = p.text
            nivel = int(p.style.name[-1])
            if nivel == 1:
                indice.add_run(f'{texto}\n').bold = True
            elif nivel == 2:
                indice.add_run(f'   {texto}\n')
                
    # Guardar el documento como un archivo de Word
    doc.save('documento.docx')

# Cerrar documento Word abierto
try:
    os.system("taskkill /f /im winword.exe")
except:
    input("Pausa")
os.system('cls')

# Crear un nuevo documento Word
doc = docx.Document("style.docx")

# Para limpiar el documento conservando la plantilla
doc._body.clear_content()

# Agregar el título principal
doc.add_heading('IV.- DESCRIPCIÓN DE LOS ELEMENTOS FISICOS Y BIOLOGICOS DE LA CUENCA HIDROLOGICO, SUBCUENCA Y MICROCUENCA DONDE SE ENCUENTRA UBICADA LA SUPERFICIE SOLICITADAS INCLUYENDO CLIMAS, TIPO DE SUELO, TOPOGRAFÍAS, HIDROGRAFÍA, GEOLOGÍA Y LA COMPOSICIÓN Y ESTRUCTURA FLORÍSTICA POR TIPOS DE VEGETACIÓN Y COMPOSICIÓN DE GRUPOS FAUNÍSTICOS.', level=1)

# Agregar el título IV.1
doc.add_heading('IV.1. Delimitación del área de estudio donde pretende establecerse el proyecto.', level=2)

# Agregar el contenido del título IV.1
doc.add_paragraph('El área sujeta al presente estudio para el establecimiento del proyecto “Extracción de Material Pétreo 2 Sabinos”, está constituida y para su análisis de este capítulo por la microcuenca “el Moral”, el cual será sujeto de análisis y en lo sucesivo se le denominará el Sistema Ambiental o SA; el cual está inmerso en la Región Hidrológica 24 “Bravo-Conchos”, en la cuenca “B” Río Bravo-Piedras Negras, específicamente en la Subcuenca “Be” Navajas. El sistema ambiental se encuentra inmerso en el municipio de Piedras Negras (100%). La cual fue considerada como una única área de estudio y para la descripción, está constituida en su mayor parte por valles con un 79.46 % y una superficie de 9,519.14 has Las características para su delimitación según la metodología que utilizó SAGARPA en su programa de FIRCO, fueron las cuencas y Subcuencas hidrológicas, fisiografía, las cartas de topografía, red hidrológica, red de caminos, altitudes, climas, poblaciones, uso de suelo y vegetación. Por su ubicación geográfica el sistema ambiental está inmerso dentro de la Región Hidrológica RH-24 Bravo Conchos, en la cuenca Río Bravo-Piedras Negras y la Subcuenca Navajas se consideró esta área ya que el área en estudio se encuentra dentro de esta. (Ver anexo Mapa 4.1.- Delimitación del Sistema Ambiental y Área de estudio).')

# Agregar el título IV.2
doc.add_heading('IV.2. Caracterización y análisis del Sistema Ambiental Hidrológico-Forestal.', level=2)

# Agregar el título IV.2.1
doc.add_heading('IV.2.1. Caracterización y análisis retrospectivo de la calidad ambiental del Sistema Ambiental.', level=3)

# Agregar el contenido del título IV.2.1

# Agregar el título IV.2.2
doc.add_heading('IV.2.2. Medio Físico', level=3)

# Agregar el título IV.2.2.1
doc.add_heading('IV.2.2.1. Clima.', level=4)

# Agregar el contenido del título IV.2.2.1
doc.add_paragraph("En el área del Sistema Ambiental se encuentra en la parte Norte del Estado de Coahuila, así mismo se encuentra inmersa en la provincia Grandes llanuras de Norteamérica. En el área que comprende el sistema ambiental se encuentran climas del Seco semicálido ya que comprende tipos de vegetación que influyen en la determinación de estos. Para representar los tipos de clima presentes, se utilizó la carta climática H14-10 (Piedras Negras), del Instituto Nacional de Estadística, Geografía e Informática (INEGI), se utilizó el conjunto de datos vectoriales del Continuo Nacional de Efectos Climáticos Regionales escala 1: 250,000, en formato digital, así como las fórmulas climáticas, se determinó de acuerdo al sistema de clasificación de Köppen modificado por Enriqueta García, encontrando que el clima más dominante es el Seco semicálido BS0hw(x') con 67.72%, siendo el semiseco semicálido BS1hx' el de menor frecuencia con 32.28 %. A continuación, se enlistan y se describen. (Ver anexo Mapa 4.2.- Tipos de climas del SA)")

# Agregar la tabla 4.1
Referencia("Tabla 4.1.- Clasificación de climas del Sistema Ambiental")
tabla = doc.add_table(rows=2, cols=5)
tabla.style = 'Table Grid'
tabla.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
# Agregar encabezados de la tabla
encabezados = tabla.rows[0].cells
encabezados[0].text = 'TIPO'
encabezados[1].text = 'CLAVE'
encabezados[2].text = 'SUPERFICIE'
encabezados[3].text = 'km2'
encabezados[4].text = 'PORCENTAJE'

# Ajustar el ancho de las columnas de la tabla
for fila in tabla.rows:
    for celda in fila.cells:
        celda.width = docx.shared.Inches(1.5)

# Agregar la descripción de los tipos de climas
doc.add_paragraph("A continuación, se describen los diferentes tipos de climas encontrados en el área del sistema ambiental a estudiar.")

# Agregar la tabla 4.2
Referencia("Tabla 4.2.- Descripción de los climas del Sistema Ambiental")
tabla = doc.add_table(rows=2, cols=3)
tabla.style = 'Table Grid'
tabla.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER

# Agregar encabezados de la tabla
encabezados = tabla.rows[0].cells
encabezados[0].text = 'Clasificación'
encabezados[1].text = 'Descripción'
encabezados[2].text = 'Vegetación de influencia'

# Ajustar el ancho de las columnas de la tabla
for fila in tabla.rows:
    for celda in fila.cells:
        celda.width = docx.shared.Inches(2.5)

# Agregar el título IV.2.2.2
doc.add_heading('IV.2.2.2. Temperatura.', level=4)

# Agregar el contenido del título IV.2.2.2
doc.add_paragraph("De acuerdo a la Estación Meteorológica de influencia en el sistema ambiental en estudio, por estar más cerca y presentar datos históricos es la estación 5025 de la Comisión Nacional del Agua (CONAGUA), ubicada en el municipio de Piedras Negras con los datos hasta el año 2010, con un histórico de 29 años, se tiene un registro de una máxima de 39 ºC en el mes de agosto, una mínima de 6.3 ºC en el mes de diciembre y una temperatura media anual de 16.3 ºC, los meses más cálidos registrados por esta estación fueron los de abril a septiembre, con temperaturas superiores a los 30ºC., los meses con temperatura más baja ocurrieron predominantemente en la época de otoño en los meses de diciembre a febrero, período durante el cual las temperaturas no superaron los 10 ºC, según los registros. Existe un registro de velocidad del viento de 13.6 km/h, donde los meses con mayor registro fueron de abril y mayo con más 15.3 km/h de acuerdo a los datos recabados por la estación del aeropuerto Piedras Negras INTL(MMPG) de las estaciones meteorológicas a este.")

# Agregar la tabla 4.3
Referencia("Tabla 4.3.- Temperaturas mínimas y máximas.")
tabla = doc.add_table(rows=2, cols=4)
tabla.style = 'Table Grid'
tabla.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER

# Agregar encabezados de la tabla
encabezados = tabla.rows[0].cells
encabezados[0].text = 'Tiempo'
encabezados[1].text = 'Temperatura Máxima (°C)'
encabezados[2].text = 'Temperatura Mínima (°C)'
encabezados[3].text = 'Temperatura Media (°C)'

# Ajustar el ancho de las columnas de la tabla
for fila in tabla.rows:
    for celda in fila.cells:
        celda.width = docx.shared.Inches(1.5)

# Agregar referencia a gráfica histórica
Referencia("Gráfica 4.1.- Temperatura histórica", grafico=True)

# Agregar el título IV.2.2.3
doc.add_heading('IV.2.2.3. Precipitación.', level=4)

# Agregar el contenido del título IV.2.2.3
doc.add_paragraph("Las precipitaciones de acuerdo a la estación meteorológica 5025 de la CONAGUA, ubicada en el municipio de Piedras Negras con los datos históricos, tiene un registro con los meses de mayor precipitación, siendo mayo y septiembre, los meses de mayor precipitación de manera oficial cuyos meses van de 64.7 a 71.2 mm, sin embargo, se tiene el registro que, para el mes de agosto, fue el mes más lluvioso, con 160 mm, sin embargo, los meses de mejor precipitación fueron febrero y diciembre con apenas 11 mm, como se puede ver en la siguiente gráfica.")

# Agregar referencia a tabla de precipitación
Referencia("Tabla 4.4.- Precipitación.")

tabla = doc.add_table(rows=2, cols=13)
tabla.style = 'Table Grid'
tabla.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER

# Agregar encabezados de la tabla
encabezados = tabla.rows[0].cells
encabezados[0].text = 'Tiemp.'
encabezados[1].text = 'Ene'
encabezados[2].text = 'Feb'
encabezados[3].text = 'Mar'
encabezados[4].text = 'Abr'
encabezados[5].text = 'May'
encabezados[6].text = 'Jun'
encabezados[7].text = 'Jul'
encabezados[8].text = 'Ago'
encabezados[9].text = 'Sep'
encabezados[10].text = 'Oct'
encabezados[11].text = 'Nov'
encabezados[12].text = 'Dic'

# Ajustar el ancho de las columnas de la tabla
for fila in tabla.rows:
    for celda in fila.cells:
        celda.width = docx.shared.Inches(0.5)

# Agregar referencia a gráfico de precipitacion
Referencia("Grafica 4.2.- Precipitación.", grafico=True)

# Agregar título
doc.add_heading("IV.2.2.3.1. Evapotranspiración.", level=4)

# Agregar contenido del título IV.2.2.4
doc.add_paragraph("Los valores mensuales de evapotranspiración se calcularon de acuerdo al método de Thornthwaite (1948), este método es basado en la determinación de la evapotranspiración en función de la temperatura media correlacionada con la duración astronómica del día y el número de días. Por lo que cuando más alta es la temperatura, mayor es el valor de evapotranspiración. En el sistema ambiental el valor de evapotranspiración acumulada es de 1855.3 mm, la mayor concentración de valores de evapotranspiración se presentó en los meses de junio a agosto, debido a que es el período de altas temperaturas, teniendo el mes de diciembre con menor evapotranspiración, de acuerdo a la estación meteorológica que registra estos datos, a continuación, se muestra la distribución de la evapotranspiración en 29 años.")

# Agregar referencia a tabla de precipitación
Referencia("Tabla 4.5.- Evapotranspiración.")

# Crear tabla
tabla = doc.add_table(rows=2, cols=13)
tabla.style = 'Table Grid'
tabla.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER

# Agregar encabezados de la tabla
encabezados = tabla.rows[0].cells
encabezados[0].text = 'Tpo. eva.'
encabezados[1].text = 'Ene'
encabezados[2].text = 'Feb'
encabezados[3].text = 'Mar'
encabezados[4].text = 'Abr'
encabezados[5].text = 'May'
encabezados[6].text = 'Jun'
encabezados[7].text = 'Jul'
encabezados[8].text = 'Ago'
encabezados[9].text = 'Sep'
encabezados[10].text = 'Oct'
encabezados[11].text = 'Nov'
encabezados[12].text = 'Dic'

# Ajustar el ancho de las columnas de la tabla
for fila in tabla.rows:
    for celda in fila.cells:
        celda.width = docx.shared.Inches(0.5)

# Referencias a gráficas
Referencia("Gráfica 4.3.- Comportamiento de evapotranspiración.", grafico=True)
Referencia("Gráfica 4.4.- Climograma.", grafico=True)

# Agregar título
doc.add_heading("IV.2.2.4. Riesgos y vulnerabilidad", level=3)

# Agregar contenido del título IV.2.2.4
doc.add_paragraph("Basado en sus características fisiográficas, geológicas y morfológicas y la ubicación geográfica del sistema ambiental, está en una zona de bajo riesgo ante la ocurrencia de diferentes fenómenos meteorológicos que pueden alterar estructuralmente las condiciones naturales, del área y el proyecto.")

# Agregar título
doc.add_heading("IV.2.2.4.1.- Riesgos Hidrometeorológicos.", level=4)

#Agregar contenido del títul
doc.add_heading("IV.2.2.4.1.1- Precipitación.", level=5)

# Contenido del título IV.2.2.4.1.1
doc.add_paragraph("El área que ocupan el sistema ambiental objeto de estudio la cual se encuentran en una zona de bajo riesgo ante la ocurrencia de este fenómeno en forma severa, ya que se encuentra, según el mapa de distribución de precipitaciones a nivel nacional, en un área de baja precipitación, por lo que su afectación no sería considerable de acuerdo a las condiciones generales del terreno. (Ver anexo Mapa 4.3.- Precipitación Media)")

# Agregar título
doc.add_heading("IV.2.2.4.1.2- Tormentas de granizo y nieve.", level=5)

# Agregar contenido del título IV.2.2.4.1.2
doc.add_paragraph("El Sistema ambiental donde se realiza el estudio por su ubicación se encuentra en un área de muy baja posibilidad de ser afectada por fenómenos, de acuerdo al mapa de Riesgo por municipio de granizadas en México del Centro Nacional de Prevención de Desastres que se muestra a continuación. (Ver anexo Mapa 4.4.- Riesgo por Granizada)")

# Agregar título
doc.add_heading("IV.2.2.4.1.3- Heladas.", level=5)

# Contenido del título IV.2.2.4.1.3
doc.add_paragraph("En el siguiente mapa de riesgos de heladas y nevadas podemos observar que el sistema ambiental en la que se pretende realizar las actividades referentes al proyecto, el área se encuentra con muy bajo a bajo de riesgo que ocurran dichos fenómenos. (Ver anexo Mapa 4.5.- Riesgo por Bajas Temperaturas)")

# Agregar título
doc.add_heading("IV.2.2.4.1.4- Ciclones tropicales.", level=5)

# Contenido del título IV.2.2.4.1.4
doc.add_paragraph("Según el mapa que se presenta, el área del sistema ambiental en la que se va a realizar el proyecto, el riesgo por el que se presente un ciclón tropical, es muy bajo, sin embargo, en los últimos años se ha presentado un aumento en la precipitación, lo que provoca una mayor humedad en general. (Ver anexo Mapa 4.6.- Riesgo por Ciclones).")

#Agregr título
doc.add_heading("IV.2.2.4.1.5- Inundaciones.",level=5)

# Contenido del título
doc.add_paragraph("El área del sistema ambiental podemos ver que se encuentra en una zona que va de baja a muy baja probabilidad de riesgo de inundaciones, de acuerdo al mapa de riesgo de inundación (Ver anexo Mapa 4.7.- Riesgo por inundación).")

# Agregar título
doc.add_heading("IV.2.2.4.1.6- Sequía.",level=5)

# Contenido del título
doc.add_paragraph("Uno de los grandes riesgos del área del sistema ambiental son las sequías, que provocan el desabasto de agua y afecta el desarrollo económico del área. Como se observa en el siguiente mapa existe riesgo alto de que se presenten sequías en los períodos de temperaturas más altas en el Estado. (Ver anexo Mapa 4.8.- Riesgo por sequias)")

# Agregar título
doc.add_heading("IV.2.2.4.1.7- Tornados.",level=5)

# Contenido del título
doc.add_paragraph("Otro de los fenómenos naturales que se han estado presentando en los últimos años son los tornados, aun cuando no se tiene está área contemplada en los Atlas de riesgo tanto de Protección Civil Estatal como del CENAPRED (Nacional), ante las condiciones extremas que se presentan y las grandes planicies, existe alta posibilidad de ocurrencia de este fenómeno, es necesario considerar a este fenómeno si bien no tornados, si fuertes vientos en el área, si las condiciones climatológicas son idóneas para este fenómeno, el proyecto requerirá de un monitoreo continuo. (Ver anexo Mapa 4.9.- Riesgo por Tornados)")

# Agregar título
doc.add_heading("IV.2.2.4.1.8- Tormentas eléctricas.",level=5)

# Contenido del título
doc.add_paragraph("Aun cuando no se tiene está área contemplada en los Atlas de riesgo tanto de Protección Civil Estatal como del CENAPRED (Nacional), como áreas de riesgo, debe de considerarse los monitoreos continuos para que no se presentes los incendios por descargas eléctricas y afecta la operación del proyecto, ya que son áreas con muy bajo riesgo. (Ver anexo Mapa 4.10.- Riesgo por tormentas eléctricas).")

# Agregar título
doc.add_heading("IV.2.2.5. Suelo.",level=3)

# Contenido del título
doc.add_paragraph("Para representar los tipos de suelos presentes en el sistema ambiental se utilizó la carta edafológica H14-10 (Piedras Negras), del Instituto Nacional de Estadística, Geografía e Informática (INEGI), se utilizó el conjunto de Datos Vectoriales del Continuo Nacional de Efectos Edafológicos escala 1: 250,000, en formato digital, encontrando lo descrito a continuación.")





















# Guardar el documento como un archivo Word
doc.save('documento.docx')
os.system("start documento.docx")